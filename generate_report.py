from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Default style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# ── Helper functions ──

def add_heading_custom(text, level=1):
    """Add heading with Arial font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False, italic=False, align=None, size=12, space_after=6, font_name='Arial'):
    """Add a paragraph with custom formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_mixed_para(parts, align=None, space_after=6):
    """Add paragraph with mixed formatting. parts = [(text, bold, italic, size, font_name), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        italic = part[2] if len(part) > 2 else False
        size = part[3] if len(part) > 3 else 12
        fn = part[4] if len(part) > 4 else 'Arial'
        run = p.add_run(text)
        run.font.name = fn
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

def add_code_block(code_text):
    """Add a code block with gray background, Courier New font."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    # Remove cell padding
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                       '<w:top w:w="30" w:type="dxa"/>'
                       '<w:left w:w="80" w:type="dxa"/>'
                       '<w:bottom w:w="30" w:type="dxa"/>'
                       '<w:right w:w="80" w:type="dxa"/>'
                       '</w:tcMar>')
    tcPr.append(tcMar)
    # Gray background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    tcPr.append(shading)
    # Set width
    table.columns[0].width = Cm(16)

    # Clear default paragraph and add code
    cell.paragraphs[0].clear()
    for line in code_text.strip().split('\n'):
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Remove first empty paragraph
    if cell.paragraphs[0].text == '':
        p = cell.paragraphs[0]
        p_element = p._element
        p_element.getparent().remove(p_element)

    return table

def add_json_block(json_text):
    """Add a JSON example with green-ish background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                       '<w:top w:w="30" w:type="dxa"/>'
                       '<w:left w:w="80" w:type="dxa"/>'
                       '<w:bottom w:w="30" w:type="dxa"/>'
                       '<w:right w:w="80" w:type="dxa"/>'
                       '</w:tcMar>')
    tcPr.append(tcMar)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F8E8" w:val="clear"/>')
    tcPr.append(shading)
    table.columns[0].width = Cm(16)

    cell.paragraphs[0].clear()
    for line in json_text.strip().split('\n'):
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2D, 0x57, 0x2D)
    if cell.paragraphs[0].text == '':
        p = cell.paragraphs[0]
        p_element = p._element
        p_element.getparent().remove(p_element)
    return table

def add_table_with_style(headers, rows, col_widths=None):
    """Add a table with header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Blue header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4" w:val="clear"/>')
        tcPr.append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            # Alternating row color
            if r_idx % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                tcPr.append(shading)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

# ══════════════════════════════════════════════════════════════
#                      COVER PAGE
# ══════════════════════════════════════════════════════════════

# Add some spacing before title
for _ in range(4):
    doc.add_paragraph()

add_para('SISTEM MANAJEMEN KONTEN BLOG / BERITA', bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('Implementasi RESTful API dengan Laravel', italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para('Laporan Praktikum', italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('Pemrograman Web Framework', italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para('Disusun oleh:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('[Nama]', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Kelas: [Kelas]', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

for _ in range(3):
    doc.add_paragraph()

add_para('Tahun 2025', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# Page break
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  KATA PENGANTAR
# ══════════════════════════════════════════════════════════════
add_heading_custom('KATA PENGANTAR', level=1)

add_para('Puji syukur ke hadirat Tuhan Yang Maha Esa atas selesainya laporan praktikum ini yang berjudul "Sistem Manajemen Konten Blog / Berita: Implementasi RESTful API dengan Laravel". Laporan ini disusun sebagai bagian dari pembelajaran mengenai pembuatan API menggunakan framework Laravel.')

add_para('Dalam laporan ini, dibahas mulai dari pengertian dasar API, pentingnya API dalam pengembangan aplikasi modern, hingga implementasi nyata menggunakan Laravel dengan studi kasus sistem manajemen konten blog dan berita.')

add_para('Penulis menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena itu, kritik dan saran yang membangun sangat diharapkan.')

add_para('', space_after=12)
add_para('Penulis', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
for _ in range(2):
    add_para('', space_after=6)
add_para('[Nama]', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  DAFTAR ISI
# ══════════════════════════════════════════════════════════════
add_heading_custom('DAFTAR ISI', level=1)

toc_items = [
    ('KATA PENGANTAR', ''),
    ('BAB 1 — APA ITU API?', ''),
    ('  1.1 Pengertian API', ''),
    ('  1.2 Cara Kerja API', ''),
    ('  1.3 Contoh API dalam Kehidupan Sehari-hari', ''),
    ('BAB 2 — KENAPA API PENTING?', ''),
    ('  2.1 Alasan API Penting', ''),
    ('  2.2 Perbandingan Tanpa API vs Dengan API', ''),
    ('BAB 3 — IMPLEMENTASI API DI LARAVEL', ''),
    ('  3.1 Route API', ''),
    ('  3.2 Studi Kasus: API Artikel', ''),
    ('  3.3 Studi Kasus: API Kategori', ''),
    ('  3.4 Studi Kasus: API Komentar', ''),
    ('  3.5 Pengambilan Data Antar Tabel (Eloquent)', ''),
    ('BAB 4 — PENUTUP', ''),
    ('  4.1 Kesimpulan', ''),
]
for item, _ in toc_items:
    add_para(item, size=12, space_after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  BAB 1 — APA ITU API?
# ══════════════════════════════════════════════════════════════
add_heading_custom('BAB 1', level=1)
add_heading_custom('APA ITU API?', level=2)

# ── 1.1 Pengertian API ──
add_heading_custom('1.1 Pengertian API', level=2)

add_para('API adalah singkatan dari Application Programming Interface. Secara sederhana, API merupakan perantara yang menghubungkan dua aplikasi agar dapat saling berkomunikasi dan bertukar data.')

add_para('Untuk memahami konsep ini, perhatikan analogi berikut ketika seseorang makan di restoran:')

add_table_with_style(
    ['Peran', 'Analogi Restoran', 'Penjelasan'],
    [
    ['Pengguna', 'Pelanggan', 'Client / Frontend'],
    ['Pramusaji', 'Pelayan', 'API (perantara)'],
    ['Dapur', 'Tempat memasak', 'Server / Database'],
    ],
    col_widths=[3, 5, 5]
)

add_para('Seorang pengguna tidak dapat langsung masuk ke dapur untuk mengambil makanan. Pengguna harus memesan terlebih dahulu melalui pramusaji. Pramusaji bertugas sebagai jembatan antara pengguna (client) dan dapur (server). Pengguna menyampaikan pesanan, pramusaji menyampaikannya ke dapur, kemudian dapur memberikan makanan melalui pramusaji kembali ke meja pengguna.')

add_para('Demikian pula cara kerja API. Aplikasi (client) mengirimkan permintaan ke server melalui API, server memproses data, kemudian API mengembalikan hasilnya ke aplikasi.')

# ── 1.2 Cara Kerja API ──
add_heading_custom('1.2 Cara Kerja API', level=2)

add_para('Cara kerja API melibatkan tiga komponen utama:')

add_mixed_para([
    ('1. Client', True, False, 12),
    (' — aplikasi yang meminta data. Contohnya aplikasi mobile, website, atau aplikasi desktop.', False, False, 12),
])

add_mixed_para([
    ('2. Server', True, False, 12),
    (' — tempat data disimpan dan diproses. Server umumnya terhubung dengan database.', False, False, 12),
])

add_mixed_para([
    ('3. API', True, False, 12),
    (' — perantara yang mengatur komunikasi antara client dan server.', False, False, 12),
])

add_para('Alur kerja API adalah sebagai berikut:')
add_para('Client → (mengirim Request) → API → (meneruskan) → Server → (memproses data) → API → (mengembalikan Response) → Client', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

add_para('')
add_para('Dengan demikian, client tidak pernah berinteraksi langsung dengan database. Semua komunikasi melalui API. Hal ini penting untuk keamanan — client tidak perlu mengetahui password database atau struktur tabel di server.')

# ── 1.3 Contoh API dalam Kehidupan Sehari-hari ──
add_heading_custom('1.3 Contoh API yang Sering Kita Pakai', level=2)

add_para('Tanpa disadari, kita menggunakan API setiap hari. Berikut beberapa contohnya:')

add_mixed_para([
    ('Login dengan Google', True, False, 12),
    (' — Aplikasi yang menyediakan tombol "Login with Google" menggunakan API milik Google. Aplikasi tersebut tidak perlu mengetahui password pengguna. Cukup Google yang melakukan verifikasi, kemudian aplikasi menerima token autentikasi.', False, False, 12),
])

add_mixed_para([
    ('Cek cuaca di ponsel', True, False, 12),
    (' — Aplikasi cuaca mengambil data dari server BMKG atau OpenWeatherMap melalui API. Setiap kali dibuka, aplikasi mengirim request ke server, kemudian server mengembalikan data suhu, kelembaban, dan prakiraan cuaca.', False, False, 12),
])

add_mixed_para([
    ('Pembayaran online', True, False, 12),
    (' — Toko online seperti Shopee atau Tokopedia tidak mengelola pembayaran sendiri. Mereka menggunakan API dari Midtrans, Xendit, atau bank. Toko hanya menerima notifikasi status pembayaran, sementara pemrosesan uang ditangani oleh penyedia payment gateway.', False, False, 12),
])

add_mixed_para([
    ('Ojek online menggunakan Google Maps', True, False, 12),
    (' — Aplikasi Gojek/Grab menampilkan peta dan rute perjalanan. Mereka tidak membuat peta sendiri, melainkan menggunakan Google Maps API.', False, False, 12),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  BAB 2 — KENAPA API PENTING?
# ══════════════════════════════════════════════════════════════
add_heading_custom('BAB 2', level=1)
add_heading_custom('KENAPA API PENTING?', level=2)

# ── 2.1 Alasan ──
add_heading_custom('2.1 Mengapa API Penting?', level=2)

add_para('Di era digital saat ini, hampir semua aplikasi modern menggunakan API. Berikut adalah alasan-alasan utamanya:')

add_para('')
add_mixed_para([
    ('1. Aplikasi mobile membutuhkan API', True, False, 12),
    (' — Aplikasi Android dan iOS tidak dapat terhubung langsung ke database server. Mereka harus berkomunikasi melalui protokol HTTP. Tanpa API, aplikasi mobile hanya menjadi aplikasi offline yang tidak dapat mengambil data secara real-time.', False, False, 12),
])

add_mixed_para([
    ('2. Frontend dan backend dapat dikembangkan secara terpisah', True, False, 12),
    (' — Tim frontend dapat mengerjakan tampilan tanpa menunggu backend selesai dengan menggunakan data dummy atau mock API. Backend juga dapat fokus mengelola data tanpa memikirkan tampilan.', False, False, 12),
])

add_mixed_para([
    ('3. Satu API melayani banyak platform', True, False, 12),
    (' — Dengan satu API, kita dapat melayani website, aplikasi Android, iOS, dan platform lainnya secara bersamaan. Tidak perlu membuat ulang backend untuk setiap platform.', False, False, 12),
])

add_mixed_para([
    ('4. Integrasi dengan pihak ketiga menjadi lebih mudah', True, False, 12),
    (' — Untuk menambahkan pembayaran online, cukup panggil API Midtrans atau Xendit. Untuk mengirim notifikasi, cukup gunakan API Firebase atau Twilio. Semua dilakukan melalui koneksi API.', False, False, 12),
])

add_mixed_para([
    ('5. Keamanan lebih terjamin', True, False, 12),
    (' — Dengan API, client tidak perlu mengetahui detail database, password, atau struktur server. Server cukup mengekspos endpoint yang diperlukan saja, sementara sisanya dilindungi.', False, False, 12),
])

# ── 2.2 Perbandingan ──
add_heading_custom('2.2 Perbandingan: Tanpa API vs Dengan API', level=2)

add_para('Untuk memperjelas perbedaannya, berikut perbandingan antara aplikasi yang menggunakan API dan yang tidak:')

add_table_with_style(
    ['Aspek', 'Tanpa API', 'Dengan API'],
    [
        ['Arsitektur', 'Frontend & backend menyatu', 'Frontend & backend terpisah'],
        ['Platform', 'Hanya dapat diakses via web', 'Web, Android, iOS, dan lainnya'],
        ['Pengembangan', 'Setiap platform memiliki backend sendiri', 'Satu backend untuk semua platform'],
        ['Maintenance', 'Perubahan UI memengaruhi backend', 'Backend tetap, UI dapat diubah kapan saja'],
        ['Integrasi pihak ketiga', 'Sulit karena tidak ada standar', 'Mudah, cukup panggil endpoint'],
        ['Skalabilitas', 'Sulit, harus melakukan scale pada semua komponen', 'Mudah, scale per service'],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

add_para('')
add_para('Kesimpulannya, API berfungsi sebagai "standar komunikasi" universal. Dengan API, aplikasi-aplikasi yang berbeda teknologi dapat saling berkomunikasi menggunakan bahasa yang sama, yaitu JSON.', italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  BAB 3 — CONTOH API DI LARAVEL
# ══════════════════════════════════════════════════════════════
add_heading_custom('BAB 3', level=1)
add_heading_custom('IMPLEMENTASI API DI LARAVEL', level=2)

add_para('Pada bab ini akan dijelaskan implementasi API menggunakan Laravel melalui studi kasus Sistem Manajemen Konten Blog dan Berita.')

add_para('Berikut adalah struktur database yang digunakan:')

add_para('')

add_table_with_style(
    ['Tabel', 'Isinya', 'Relasi'],
    [
        ['users', 'Data pengguna (id, name, email, password, role, avatar)', 'Memiliki articles & comments'],
        ['categories', 'Kategori konten (id, name, slug)', 'Digunakan oleh articles'],
        ['articles', 'Artikel/konten berita (id, title, slug, content, status, dll)', 'Memiliki category, user, tags, comments'],
        ['comments', 'Komentar pembaca (id, article_id, user_id, content, status)', 'Memiliki article & user'],
        ['tags', 'Label topik (id, name, slug)', 'Many-to-many dengan articles'],
        ['article_tags', 'Tabel pivot — menghubungkan article & tag', '—'],
    ],
    col_widths=[2.5, 6.5, 5.5]
)

add_para('')
add_para('Struktur tabel di atas mencakup relasi one-to-many (category → articles), many-to-many (articles ↔ tags), dan relasi bertingkat (article → comments → user). Semua relasi ini akan diakses menggunakan Eloquent ORM Laravel, yang memungkinkan penulisan query menjadi lebih singkat dan mudah dibaca.', italic=True)

# ── 3.1 Route API ──
add_heading_custom('3.1 Route API', level=2)

add_para('Semua route API didefinisikan di dalam file routes/api.php. Terdapat tiga kelompok route: publik, memerlukan autentikasi, dan admin. Berikut daftar lengkapnya:')

add_table_with_style(
    ['Method', 'Endpoint', 'Deskripsi', 'Auth'],
    [
        ['POST', '/api/auth/register', 'Daftar akun baru', 'Publik'],
        ['POST', '/api/auth/login', 'Login dan mendapatkan token', 'Publik'],
        ['POST', '/api/auth/logout', 'Hapus token (logout)', 'Sanctum'],
        ['GET', '/api/auth/me', 'Lihat data pengguna yang login', 'Sanctum'],
        ['GET', '/api/articles', 'Ambil daftar artikel', 'Publik'],
        ['GET', '/api/articles/{slug}', 'Ambil detail artikel', 'Publik'],
        ['POST', '/api/articles', 'Buat artikel baru', 'Sanctum'],
        ['PUT', '/api/articles/{id}', 'Perbarui artikel', 'Sanctum'],
        ['DELETE', '/api/articles/{id}', 'Hapus artikel', 'Sanctum'],
        ['POST', '/api/articles/{id}/publish', 'Publikasikan artikel', 'Sanctum'],
        ['GET', '/api/articles/{id}/comments', 'Ambil komentar artikel', 'Publik'],
        ['POST', '/api/comments', 'Tambah komentar', 'Sanctum'],
        ['PUT', '/api/comments/{id}', 'Perbarui komentar', 'Sanctum'],
        ['DELETE', '/api/comments/{id}', 'Hapus komentar', 'Sanctum'],
        ['GET', '/api/categories', 'Ambil daftar kategori', 'Publik'],
        ['GET', '/api/categories/{slug}/articles', 'Ambil artikel per kategori', 'Publik'],
        ['GET', '/api/tags', 'Ambil daftar tag', 'Publik'],
        ['GET', '/api/tags/{slug}/articles', 'Ambil artikel per tag', 'Publik'],
        ['GET', '/api/admin/comments', 'Semua komentar (khusus admin)', 'Admin'],
        ['PUT', '/api/admin/comments/{id}/approve', 'Setujui komentar', 'Admin'],
    ],
    col_widths=[1.5, 4.5, 5.5, 2]
)

add_para('')
add_para('Berikut adalah contoh potongan kode dari routes/api.php:', space_after=4)

add_code_block('''// Route publik
Route::post('/auth/register', [AuthController::class, 'register']);
Route::get('/articles', [ArticleController::class, 'index']);

// Route butuh login (Sanctum)
Route::middleware('auth:sanctum')->group(function () {
    Route::post('/articles', [ArticleController::class, 'store']);
    Route::post('/articles/{id}/publish', [ArticleController::class, 'publish']);
});

// Route khusus admin
Route::middleware(['auth:sanctum', 'admin'])->group(function () {
    Route::get('/admin/comments', [CommentController::class, 'adminIndex']);
});''')

add_para('')
add_para('Perhatikan bahwa route publik dapat diakses oleh siapa saja, route Sanctum memerlukan token autentikasi, dan route admin hanya dapat diakses oleh pengguna dengan role admin. Pengecekan admin menggunakan custom middleware yang memeriksa role pengguna.', italic=True)

# ── 3.2 Studi Kasus 1 — API Artikel ──
add_heading_custom('3.2 Studi Kasus 1 — API Artikel', level=2)

add_para('Fitur utama dari API ini adalah pengambilan data artikel. Terdapat dua method penting di ArticleController, yaitu index() untuk menampilkan daftar artikel dan show() untuk menampilkan detail artikel.')

add_para('')
add_para('Method: index() — Mengambil Seluruh Artikel', bold=True, size=12)

add_para('Method index() akan mengambil artikel yang berstatus "published", lengkap dengan data kategori, penulis, dan tag. Data juga dapat difilter berdasarkan kategori, tag, penulis, atau pencarian judul/konten.', space_after=4)

add_code_block('''public function index(Request $request): JsonResponse
{
    $query = Article::with(['category', 'user', 'tags'])
        ->withCount('comments')
        ->where('status', 'published');

    // Filter berdasarkan kategori
    if ($request->filled('category')) {
        $query->whereHas('category', fn($q) =>
            $q->where('slug', $request->category));
    }

    // Pencarian judul/konten
    if ($request->filled('search')) {
        $search = $request->search;
        $query->where(function ($q) use ($search) {
            $q->where('title', 'like', "%{$search}%")
              ->orWhere('content', 'like', "%{$search}%");
        });
    }

    $articles = $query->latest('published_at')->paginate(10);

    return $this->success(
        $articles->items(),
        'Data artikel berhasil diambil',
        200,
        [
            'current_page' => $articles->currentPage(),
            'total'        => $articles->total(),
            'per_page'     => $articles->perPage(),
            'last_page'    => $articles->lastPage(),
        ]
    );
}''')

add_para('')
add_para('Berikut adalah contoh hasil JSON dari method index():', space_after=4)

add_json_block('''{
    "success": true,
    "message": "Data artikel berhasil diambil",
    "data": [
        {
            "id": 1,
            "title": "Laravel 11 Resmi Dirilis",
            "slug": "laravel-11-resmi-dirilis-abc12",
            "category": { "id": 1, "name": "Teknologi", "slug": "teknologi" },
            "user": { "id": 1, "name": "Admin", "email": "admin@blog.test" },
            "tags": [
                { "id": 1, "name": "laravel", "slug": "laravel" }
            ],
            "comments_count": 5,
            "views": 1200,
            "published_at": "2025-12-01T08:00:00.000000Z"
        }
    ],
    "meta": {
        "current_page": 1,
        "total": 25,
        "per_page": 10,
        "last_page": 3
    }
}''')

add_para('')
add_para('Dengan menggunakan with([...category, user, tags]), Laravel mengambil data dari 4 tabel sekaligus dalam satu kali eksekusi. withCount(comments) menambahkan kolom comments_count tanpa perlu query manual. Filter menggunakan whereHas untuk mencari berdasarkan relasi — misalnya mencari artikel yang termasuk dalam kategori \"teknologi\".', italic=True)

add_para('')
add_para('Method: show() — Detail Artikel', bold=True, size=12)

add_para('Method ini mengambil satu artikel berdasarkan slug (bukan ID), lengkap dengan komentar yang telah disetujui beserta data pengguna yang berkomentar:', space_after=4)

add_code_block('''public function show(string $slug): JsonResponse
{
    $article = Article::with([
        'category', 'user', 'tags',
        'comments' => function ($q) {
            $q->where('status', 'approved')->with('user');
        }
    ])->where('slug', $slug)->firstOrFail();

    $article->increment('views');

    return $this->success($article, 'Detail artikel berhasil diambil');
}''')

add_para('')
add_para('Hasil JSON dari show() — detail artikel + komentar:')

add_json_block('''{
    "success": true,
    "message": "Detail artikel berhasil diambil",
    "data": {
        "id": 1,
        "title": "Laravel 11 Resmi Dirilis",
        "content": "<p>Laravel 11 hadir dengan fitur...</p>",
        "comments": [
            {
                "id": 1,
                "content": "Mantap!",
                "status": "approved",
                "user": {
                    "id": 3,
                    "name": "Budi Santoso"
                }
            }
        ]
    }
}''')

add_para('')
add_para('Fungsi firstOrFail() secara otomatis akan mengembalikan error 404 apabila slug tidak ditemukan. Increment(views) akan menambahkan 1 ke kolom views setiap kali artikel dilihat.', italic=True)

doc.add_page_break()

# ── 3.3 Studi Kasus 2 — API Kategori ──
add_heading_custom('3.3 Studi Kasus 2 — API Kategori', level=2)

add_para('Kategori memiliki relasi one-to-many dengan artikel, artinya satu kategori dapat digunakan oleh banyak artikel. Method index() di CategoryController mengambil semua kategori beserta jumlah artikel di setiap kategori:', space_after=4)

add_code_block('''public function index(): JsonResponse
{
    $categories = Category::withCount('articles')->get();

    return $this->success(
        $categories,
        'Data kategori berhasil diambil'
    );
}''')

add_para('')
add_para('Hasil JSON-nya:', space_after=4)

add_json_block('''{
    "success": true,
    "message": "Data kategori berhasil diambil",
    "data": [
        {
            "id": 1,
            "name": "Teknologi",
            "slug": "teknologi",
            "articles_count": 15
        },
        {
            "id": 2,
            "name": "Politik",
            "slug": "politik",
            "articles_count": 8
        },
        {
            "id": 3,
            "name": "Olahraga",
            "slug": "olahraga",
            "articles_count": 12
        }
    ]
}''')

add_para('')
add_para('Kode di atas hanya terdiri dari 2 baris — ketika menggunakan withCount(articles), Laravel secara otomatis menambahkan kolom articles_count melalui query COUNT di database. Sangat efisien.', italic=True)

add_para('')
add_para('Selain itu, terdapat method articlesByCategory() untuk mengambil artikel-artikel dalam satu kategori tertentu:', space_after=4)

add_code_block('''public function articlesByCategory(string $slug): JsonResponse
{
    $category = Category::where('slug', $slug)->firstOrFail();
    $articles = Article::with(['user', 'tags'])
        ->where('category_id', $category->id)
        ->where('status', 'published')
        ->latest('published_at')
        ->paginate(10);

    return $this->success(
        $articles->items(),
        "Artikel kategori {$category->name} berhasil diambil",
        200,
        [
            'current_page' => $articles->currentPage(),
            'total'        => $articles->total(),
            'per_page'     => $articles->perPage(),
            'last_page'    => $articles->lastPage(),
        ]
    );
}''')

# ── 3.4 Studi Kasus 3 — API Komentar ──
add_heading_custom('3.4 Studi Kasus 3 — API Komentar', level=2)

add_para('Dalam sistem ini, komentar memiliki status: pending (menunggu persetujuan), approved (disetujui), atau rejected (ditolak). Pengguna biasa hanya dapat melihat komentar yang berstatus approved. Admin dapat menyetujui komentar melalui endpoint terpisah.', space_after=4)

add_code_block('''// Ambil komentar publik (hanya yang sudah disetujui)
public function index(string $articleId): JsonResponse
{
    $comments = Comment::where('article_id', $articleId)
        ->where('status', 'approved')
        ->with('user')
        ->latest()
        ->get();

    return $this->success($comments, 'Komentar berhasil diambil');
}

// Tambah komentar (langsung status pending)
public function store(StoreCommentRequest $request): JsonResponse
{
    $data = $request->validated();
    $data['user_id'] = $request->user()->id;
    $data['status'] = 'pending';

    $comment = Comment::create($data);
    $comment->load('user');

    return $this->success(
        $comment,
        'Komentar berhasil ditambahkan (menunggu approve)',
        201
    );
}''')

add_para('')
add_para('Hasil JSON komentar:', space_after=4)

add_json_block('''{
    "success": true,
    "message": "Komentar berhasil diambil",
    "data": [
        {
            "id": 1,
            "content": "Artikelnya sangat bermanfaat!",
            "status": "approved",
            "created_at": "2 jam lalu",
            "user": {
                "id": 3,
                "name": "Budi Santoso",
                "avatar": null
            }
        },
        {
            "id": 2,
            "content": "Terima kasih, ditunggu artikel selanjutnya.",
            "status": "approved",
            "created_at": "1 jam lalu",
            "user": {
                "id": 5,
                "name": "Siti Rahmawati",
                "avatar": null
            }
        }
    ]
}''')

add_para('')
add_para('Dengan menggunakan with(user) pada query komentar, Laravel akan mengambil nama, email, dan avatar dari pengguna yang menulis komentar — semuanya dalam satu query. Tidak perlu melakukan looping manual atau menulis JOIN.', italic=True)

add_para('')
add_para('Sementara itu, terdapat endpoint khusus untuk admin guna melihat seluruh komentar (termasuk yang masih pending) dan menyetujui komentar:', space_after=4)

add_code_block('''// Khusus admin: lihat seluruh komentar
public function adminIndex(): JsonResponse
{
    $comments = Comment::with(['article', 'user'])
        ->latest()->get();

    return $this->success(
        $comments,
        'Semua komentar berhasil diambil'
    );
}

// Khusus admin: approve komentar
public function approve(string $id): JsonResponse
{
    $comment = Comment::findOrFail($id);
    $comment->update(['status' => 'approved']);

    return $this->success(
        $comment,
        'Komentar berhasil diapprove'
    );
}''')

doc.add_page_break()

# ── 3.5 Pengambilan Data Antar Tabel ──
add_heading_custom('3.5 Pengambilan Data Antar Tabel (Eloquent Relationship)', level=2)

add_para('Salah satu keunggulan Laravel adalah Eloquent ORM. Dengan Eloquent, kita dapat mengambil data dari beberapa tabel sekaligus tanpa perlu menulis query JOIN secara manual. Cukup menggunakan method relasi di model, dan Laravel akan menangani sisanya.')

add_para('')
add_para('Relasi 1: Article → Category (belongsTo)', bold=True, size=12)

add_para('Setiap artikel memiliki satu kategori. Berikut adalah definisi relasi di model Article:', space_after=4)

add_code_block('''// Model Article
public function category()
{
    return $this->belongsTo(Category::class);
}

// Cara panggil di controller
$article = Article::with('category')->first();''')

add_para('')
add_para('Hasil JSON:', space_after=4)

add_json_block('''{
    "id": 1,
    "title": "Laravel 11 Resmi Dirilis",
    "category": {
        "id": 1,
        "name": "Teknologi",
        "slug": "teknologi"
    }
}''')

add_para('')
add_para('Laravel secara otomatis mencocokkan category_id di tabel articles dengan id di tabel categories. Hasilnya langsung digabung menjadi satu objek JSON. ', italic=True)
add_para('Tanpa with(category), kita hanya mendapatkan angka category_id, bukan data kategorinya.', italic=True)

add_para('')
add_para('Relasi 2: Category → Articles (hasMany)', bold=True, size=12)

add_para('Ini adalah kebalikan dari relasi sebelumnya: satu kategori dapat digunakan oleh banyak artikel:', space_after=4)

add_code_block('''// Model Category
public function articles()
{
    return $this->hasMany(Article::class);
}

// Cara panggil di controller
$category = Category::withCount('articles')->first();''')

add_para('')
add_para('Hasil JSON:', space_after=4)

add_json_block('''{
    "id": 1,
    "name": "Teknologi",
    "slug": "teknologi",
    "articles_count": 15
}''')

add_para('')
add_para('withCount(articles) menambahkan kolom articles_count yang nilainya dihitung secara otomatis dari database. Berguna untuk menampilkan jumlah artikel di setiap kategori tanpa perlu menghitung manual.', italic=True)

add_para('')
add_para('Relasi 3: Article → Comments → User (Nested)', bold=True, size=12)

add_para('Berikut adalah contoh relasi bertingkat: artikel memiliki banyak komentar, dan setiap komentar ditulis oleh satu pengguna:', space_after=4)

add_code_block('''// Model Article
public function comments()
{
    return $this->hasMany(Comment::class);
}

// Model Comment
public function user()
{
    return $this->belongsTo(User::class);
}

// Ambil artikel + komentar + user — 3 level!
$article = Article::with(['comments' => function ($q) {
    $q->where('status', 'approved')->with('user');
}])->first();''')

add_para('')
add_para('Hasil JSON:', space_after=4)

add_json_block('''{
    "id": 1,
    "title": "Laravel 11 Resmi Dirilis",
    "comments": [
        {
            "id": 1,
            "content": "Mantap!",
            "status": "approved",
            "user": {
                "id": 3,
                "name": "Budi Santoso",
                "email": "budi@example.com"
            }
        }
    ]
}''')

add_para('')
add_para('Ini yang disebut dengan eager loading. Tanpa with(), setiap artikel akan menjalankan query tambahan untuk mengambil komentar (masalah N+1). Dengan with(), semua data dari 3 tabel diambil hanya dalam 3 query — sangat efisien.', italic=True)

add_para('')
add_para('Relasi 4: Article ↔ Tags (Many-to-Many via Pivot)', bold=True, size=12)

add_para('Satu artikel dapat memiliki banyak tag, dan satu tag dapat digunakan oleh banyak artikel. Hubungan ini disimpan di tabel pivot article_tags:', space_after=4)

add_code_block('''// Model Article
public function tags()
{
    return $this->belongsToMany(Tag::class, 'article_tags');
}

// Model Tag
public function articles()
{
    return $this->belongsToMany(Article::class, 'article_tags');
}

// Cara panggil — semudah relasi lainnya
$article = Article::with('tags')->first();''')

add_para('')
add_para('Hasil JSON:', space_after=4)

add_json_block('''{
    "id": 1,
    "title": "Laravel 11 Resmi Dirilis",
    "tags": [
        { "id": 1, "name": "laravel", "slug": "laravel" },
        { "id": 3, "name": "php", "slug": "php" }
    ]
}''')

add_para('')
add_para('Laravel secara otomatis mendeteksi tabel pivot article_tags, mencocokkan article_id dengan tag_id, dan mengembalikan data tag. Cukup dengan menulis belongsToMany, Laravel akan menangani sisanya.', italic=True)

add_para('')
add_para('Tabel ringkasan seluruh relasi yang digunakan:', space_after=4)

add_table_with_style(
    ['Tipe', 'Model 1', 'Model 2', 'Penjelasan'],
    [
        ['belongsTo', 'Article', 'Category', 'Artikel memiliki 1 kategori'],
        ['hasMany', 'Category', 'Article', 'Kategori memiliki banyak artikel'],
        ['belongsTo', 'Comment', 'Article', 'Komentar milik 1 artikel'],
        ['belongsTo', 'Comment', 'User', 'Komentar ditulis oleh 1 user'],
        ['hasMany', 'Article', 'Comment', 'Artikel memiliki banyak komentar'],
        ['belongsToMany', 'Article', 'Tag', 'Artikel banyak tag, tag banyak artikel (pivot)'],
    ],
    col_widths=[2.5, 2.5, 2.5, 7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  BAB 4 — PENUTUP
# ══════════════════════════════════════════════════════════════
add_heading_custom('BAB 4', level=1)
add_heading_custom('PENUTUP', level=2)

add_heading_custom('4.1 Kesimpulan', level=2)

add_para('Setelah mempelajari API dan implementasinya di Laravel, dapat ditarik beberapa kesimpulan penting sebagai berikut:')

add_para('')
add_mixed_para([
    ('1. API merupakan penghubung antara aplikasi dan data.', True, False, 12),
    (' API bertindak sebagai perantara yang mengatur komunikasi antara client (aplikasi web, mobile, dll) dan server (database). Tanpa API, aplikasi modern tidak dapat berfungsi secara optimal, terutama aplikasi mobile yang memerlukan data dari internet.', False, False, 12),
])

add_mixed_para([
    ('2. Laravel mempermudah pembuatan API melalui fitur Eloquent ORM yang powerful.', True, False, 12),
    (' Dengan beberapa baris kode, kita dapat mengambil data dari beberapa tabel sekaligus. Fitur eager loading (with()), pagination, dan query builder menjadikan pengembangan API lebih cepat dan terstruktur.', False, False, 12),
])

add_mixed_para([
    ('3. Relasi antar tabel dapat diakses sekaligus tanpa query berulang.', True, False, 12),
    (' Melalui Eloquent relationship (belongsTo, hasMany, belongsToMany), kita dapat mengambil data artikel, kategori, penulis, tag, komentar, dan pengguna — semua dalam jumlah query yang minimal. Efisien, mudah dibaca, dan tidak perlu menulis JOIN secara manual.', False, False, 12),
])

add_para('')
add_para('Demikian laporan ini disusun. Semoga dapat memberikan pemahaman yang jelas mengenai konsep API dan implementasinya menggunakan Laravel.', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Belajar_API_di_Laravel.docx')
doc.save(output_path)
print(f'Dokumen berhasil dibuat: {output_path}')
